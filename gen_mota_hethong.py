# -*- coding: utf-8 -*-
"""Sinh file Word mô tả Hệ thống Quản lý Bếp ăn Thông minh."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ORANGE = RGBColor(0xD2, 0x69, 0x1E)
DARK = RGBColor(0x20, 0x20, 0x20)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ---- Style mặc định ----
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(11)
normal.font.color.rgb = DARK


def set_heading_color(p, color):
    for run in p.runs:
        run.font.color.rgb = color


def h1(text):
    p = doc.add_heading(text, level=1)
    set_heading_color(p, ORANGE)
    return p


def h2(text):
    p = doc.add_heading(text, level=2)
    set_heading_color(p, RGBColor(0x8B, 0x45, 0x13))
    return p


def para(text, italic=False, bold=False, size=11, color=DARK, align=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    r.font.size = Pt(size)
    r.font.color.rgb = color
    if align is not None:
        p.alignment = align
    return p


def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# =========================================================
# TRANG BÌA
# =========================================================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('TÀI LIỆU MÔ TẢ HỆ THỐNG')
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = ORANGE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('HỆ THỐNG QUẢN LÝ BẾP ĂN THÔNG MINH')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = DARK

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('(Smart Kitchen Management System)')
r.italic = True
r.font.size = Pt(14)
r.font.color.rgb = GREY

doc.add_paragraph()
unit = doc.add_paragraph()
unit.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = unit.add_run('Đơn vị: KV2 - VT NET')
r.font.size = Pt(13)
r.bold = True

dom = doc.add_paragraph()
dom.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = dom.add_run('Tên miền: net2kitchen.viettel.pro.vn')
r.font.size = Pt(12)
r.font.color.rgb = GREY

doc.add_paragraph()
doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = info.add_run('Nền tảng: Django 5.x · PostgreSQL · Google Gemini AI · NetChat Bot\n'
                 'Triển khai: Ubuntu Server · Nginx · Gunicorn')
r.font.size = Pt(11)
r.font.color.rgb = GREY

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = date_p.add_run('\nNgày cập nhật: 22/06/2026')
r.italic = True
r.font.size = Pt(11)

doc.add_page_break()

# =========================================================
# 1. GIỚI THIỆU CHUNG
# =========================================================
h1('1. Giới thiệu chung')
para('Hệ thống Quản lý Bếp ăn Thông minh là giải pháp phần mềm toàn diện nhằm tự '
     'động hóa toàn bộ quy trình vận hành bếp ăn nội bộ của đơn vị KV2 - VT NET. '
     'Hệ thống số hóa các công đoạn từ đăng ký suất ăn, lên thực đơn, tính toán '
     'nguyên liệu, quản lý thu - chi tài chính, cho đến điểm danh thực tế bằng '
     'camera nhận diện khuôn mặt và đánh giá chất lượng bữa ăn.')
para('Điểm khác biệt của hệ thống là việc tích hợp Trí tuệ nhân tạo (Google Gemini) '
     'để phân tích dinh dưỡng, gợi ý thực đơn cân bằng; kết hợp hệ thống thông báo '
     'tự động qua Bot NetChat (tương thích Mattermost) giúp tối ưu chi phí, minh '
     'bạch tài chính và nâng cao trải nghiệm cho người ăn.')

h2('1.1. Mục tiêu')
bullet('giảm thiểu sai sót khi tổng hợp danh sách đăng ký ăn hàng ngày.', 'Tự động hóa: ')
bullet('mua nguyên liệu vừa đủ theo số người đăng ký, giảm lãng phí.', 'Tiết kiệm: ')
bullet('công khai thu - chi, có quy trình phê duyệt nhiều cấp.', 'Minh bạch: ')
bullet('thực đơn cân bằng dinh dưỡng, có lời khuyên sức khỏe từ AI.', 'Sức khỏe: ')
bullet('điểm danh nhanh bằng nhận diện khuôn mặt, chúc mừng sinh nhật, đánh giá món ăn.', 'Gắn kết: ')

h2('1.2. Đối tượng sử dụng')
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Light Grid Accent 2'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
hdr[0].paragraphs[0].add_run('Vai trò').bold = True
hdr[1].paragraphs[0].add_run('Chức năng chính').bold = True
roles = [
    ('Admin (Quản trị)', 'Phê duyệt thực đơn/món/chi phí, cấu hình hệ thống (Bot, AI, giá suất ăn), quản lý người dùng, import danh sách.'),
    ('Nhân viên bếp', 'Tạo thực đơn, khai báo món ăn & định mức nguyên liệu, nhập đơn mua hàng, quản lý kho và chi phí.'),
    ('Người ăn (Diner)', 'Đăng nhập OTP, xem thực đơn, chuyển suất ăn, đánh giá bữa ăn, đề xuất món mới.'),
]
for role, func in roles:
    cells = tbl.add_row().cells
    cells[0].paragraphs[0].add_run(role).bold = True
    cells[1].text = func

# =========================================================
# 2. CÔNG NGHỆ & KIẾN TRÚC
# =========================================================
h1('2. Công nghệ & Kiến trúc hệ thống')

h2('2.1. Công nghệ sử dụng')
tech = [
    ('Backend', 'Python 3.12, Django 5.x (kiến trúc MVT)'),
    ('Cơ sở dữ liệu', 'PostgreSQL'),
    ('Frontend', 'HTML5, CSS3 (Vanilla + Bootstrap 5), JavaScript (ES6+), Chart.js'),
    ('Trí tuệ nhân tạo', 'Google Generative AI - Gemini (phân tích dinh dưỡng, gợi ý thực đơn)'),
    ('Giao tiếp / Thông báo', 'NetChat API (tương thích Mattermost) - gửi OTP & nhắc nhở'),
    ('Nhận diện khuôn mặt', 'Client camera (thư mục Client/) gửi heartbeat & dữ liệu điểm danh qua API'),
    ('Triển khai', 'Ubuntu Server, Nginx (reverse proxy), Gunicorn (WSGI)'),
]
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Light Grid Accent 2'
hdr = tbl.rows[0].cells
hdr[0].paragraphs[0].add_run('Hạng mục').bold = True
hdr[1].paragraphs[0].add_run('Công nghệ').bold = True
for k, v in tech:
    cells = tbl.add_row().cells
    cells[0].paragraphs[0].add_run(k).bold = True
    cells[1].text = v

h2('2.2. Cấu trúc module (Django apps)')
mods = [
    ('accounts', 'Quản lý người dùng, hồ sơ (UserProfile), phân quyền (RBAC) và đăng nhập OTP.'),
    ('core', 'Tính năng lõi: Dashboard, dịch vụ AI, điểm danh (AttendanceLog), heartbeat & ảnh chụp camera, màn hình TV Show, chúc mừng sinh nhật, cấu hình hệ thống.'),
    ('meals', 'Quản lý món ăn, nguyên liệu, định mức, thực đơn ngày/tuần, gợi ý AI và đơn chuẩn bị nguyên liệu.'),
    ('finance', 'Quản lý thu - chi: đơn mua hàng, chi phí, hóa đơn, tồn kho, cấu hình giá suất ăn, phân tích dinh dưỡng.'),
    ('registrations', 'Quản lý đăng ký suất ăn, chuyển suất ăn, log gửi thông báo.'),
    ('reviews', 'Đánh giá bữa ăn / món ăn (qua QR), đề xuất món mới và bình chọn.'),
]
tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Light Grid Accent 2'
hdr = tbl.rows[0].cells
hdr[0].paragraphs[0].add_run('Module').bold = True
hdr[1].paragraphs[0].add_run('Vai trò').bold = True
for k, v in mods:
    cells = tbl.add_row().cells
    cells[0].paragraphs[0].add_run(k).bold = True
    cells[1].text = v

para('Ngoài ra, hệ thống còn có các thư mục: templates/ (giao diện), static/ '
     '(CSS, JS, ảnh), media/ (ảnh món ăn, avatar, ảnh hóa đơn, ảnh chụp nhận diện) '
     'và Client/ (chương trình camera nhận diện khuôn mặt đặt tại bếp).', italic=True, color=GREY)

# =========================================================
# 3. CÁC PHÂN HỆ CHỨC NĂNG
# =========================================================
h1('3. Các phân hệ chức năng')

h2('3.1. Dashboard & Phân tích AI')
bullet('Tổng quan thời gian thực: số người đăng ký, chi phí thu - chi và chênh lệch tài chính trong ngày.')
bullet('AI Dinh dưỡng: tự động phân tích hàm lượng calo và đưa ra lời khuyên sức khỏe dựa trên thực đơn thực tế qua Gemini (lưu tại DailyNutritionAnalysis).')
bullet('Xu hướng tài chính: biểu đồ so sánh thu - chi theo tuần / tháng / năm bằng Chart.js.')

h2('3.2. Quản lý Đăng ký & Điểm danh')
bullet('Import đăng ký: nhập danh sách đăng ký ăn hàng ngày từ file Excel (upsert an toàn - chỉ cập nhật ô có dữ liệu, ô trống giữ nguyên).')
bullet('Gửi thông báo Bot: tự động gửi tin nhắn nhắc nhở / xác nhận qua NetChat tới từng nhân viên; ghi log kết quả (NotificationLog).')
bullet('Chuyển suất ăn: người dùng có thể chuyển suất ăn của mình ngày X cho người khác (trước 11h); hệ thống tự áp dụng khi dữ liệu được đồng bộ.')
bullet('Điểm danh thực tế: tab "Tham gia" hiển thị danh sách người ăn thực tế, đồng bộ qua API từ hệ thống camera nhận diện (AttendanceLog).')

h2('3.3. Quản lý Thực đơn & Món ăn')
bullet('Kho món ăn: quản lý chi tiết món (loại món chính/phụ/canh/tráng miệng, khẩu phần, hình ảnh, mô tả).')
bullet('Công thức & nguyên liệu: khai báo định mức nguyên liệu cho từng món để tự động tính khối lượng cần mua theo số người đăng ký.')
bullet('AI gợi ý thực đơn: AI đề xuất thực đơn 5 ngày trong tuần dựa trên tiêu chí cân bằng dinh dưỡng và đa dạng (WeeklyMenuDraft).')
bullet('Quy trình phê duyệt: thực đơn và món mới do bếp tạo ở trạng thái "Chờ duyệt" → Admin duyệt / từ chối (có ghi log lý do) trước khi áp dụng.')
bullet('Đơn chuẩn bị nguyên liệu: sau khi xác nhận, hệ thống sinh đơn nguyên liệu (MenuPrepOrder) làm cơ sở cho hóa đơn đặt hàng.')

h2('3.4. Quản lý Kho & Tài chính')
bullet('Đặt hàng: tạo đơn mua nguyên liệu chính (main) và đơn mua bổ sung / gia vị (extra), kèm ảnh bill.')
bullet('Quản lý chi phí: ghi nhận, phân loại mọi chi phí; quy trình phê duyệt minh bạch, có lịch sử chỉnh sửa và từ chối.')
bullet('Quản lý tồn kho: theo dõi nguyên liệu còn lại theo từng ngày, nhập tay hoặc trích từ hóa đơn; lưu lịch sử nhập / xuất kho.')
bullet('Cấu hình giá: thiết lập giá suất ăn và giá gia vị theo từng giai đoạn thời gian, có log thay đổi giá.')

h2('3.5. Đánh giá & Tương tác')
bullet('Đánh giá bữa ăn qua QR Code: người ăn quét mã, chấm điểm từng món (1-5 sao) và góp ý; hỗ trợ đánh giá ẩn danh theo session.')
bullet('Đề xuất món ăn: người dùng đề xuất món mới muốn bếp nấu và bình chọn (mỗi người 1 phiếu / món).')

h2('3.6. Màn hình TV Show tại bếp')
bullet('Trình chiếu thực đơn trong ngày trên màn hình TV (Samsung) đặt tại khu vực bếp ăn.')
bullet('Chúc mừng sinh nhật: khi nhân viên có sinh nhật điểm danh, sau 5 phút màn hình hiển thị lời chúc mừng (chỉ chiếu 1 lần / ngày).')
bullet('Cảnh báo camera offline: theo dõi heartbeat camera; nếu mất tín hiệu sau 11h03 sẽ gửi cảnh báo NetChat, nhắc lại mỗi 2 phút.')

# =========================================================
# 4. PHÂN QUYỀN & ĐĂNG NHẬP OTP
# =========================================================
h1('4. Phân quyền & Bảo mật')
h2('4.1. Phân quyền (RBAC)')
para('Hệ thống chia 3 vai trò: Admin, Nhân viên bếp và Người ăn. Mỗi vai trò có '
     'phạm vi truy cập và thao tác riêng, đảm bảo nguyên tắc tối thiểu quyền.')
h2('4.2. Đăng nhập bằng OTP')
bullet('Người dùng đăng nhập bằng mã OTP (6 số) nhận qua NetChat thay cho mật khẩu.')
bullet('OTP có hiệu lực 10 phút, dùng một lần (OTPToken).')
bullet('Chống brute-force: khóa theo mã nhân viên - tối đa 10 lần sai sẽ khóa 15 phút; giới hạn gửi lại OTP mỗi 60 giây.')

# =========================================================
# 5. TÍCH HỢP
# =========================================================
h1('5. Tích hợp bên ngoài')
h2('5.1. Trí tuệ nhân tạo Google Gemini')
bullet('Phân tích dinh dưỡng thực đơn (tổng calo, mức độ, lời khuyên).')
bullet('Gợi ý thực đơn tuần cân bằng và đa dạng.')
bullet('Model và API key cấu hình linh hoạt trong Profile Admin (runtime).')

h2('5.2. Bot NetChat (tương thích Mattermost)')
bullet('Gửi mã OTP đăng nhập.')
bullet('Nhắc nhở / xác nhận đăng ký suất ăn.')
bullet('Cảnh báo camera nhận diện offline.')
bullet('URL, Token và mẫu tin nhắn được cấu hình trong Profile Admin.')

h2('5.3. Hệ thống camera nhận diện khuôn mặt')
bullet('Chương trình client (thư mục Client/) chạy tại bếp, nhận diện khuôn mặt nhân viên.')
bullet('Gửi heartbeat định kỳ để báo trạng thái sống (RecognitionHeartbeat).')
bullet('Đẩy dữ liệu điểm danh (AttendanceLog) và ảnh chụp khung hình (AttendanceCapture, tự xóa sau 30 ngày) về server.')
bullet('Server đối soát người điểm danh với danh sách đăng ký để xác định "Đã đăng ký" / "Chưa đăng ký".')

# =========================================================
# 6. MÔ HÌNH DỮ LIỆU
# =========================================================
h1('6. Mô hình dữ liệu chính')
data = [
    ('UserProfile / OTPToken', 'accounts', 'Hồ sơ người dùng, vai trò, avatar; mã OTP đăng nhập.'),
    ('Dish / Ingredient / DishIngredient', 'meals', 'Món ăn, nguyên liệu và định mức nguyên liệu theo món.'),
    ('DailyMenu / DailyMenuItem', 'meals', 'Thực đơn theo ngày và các món trong thực đơn.'),
    ('WeeklyMenuDraft', 'meals', 'Bản nháp thực đơn tuần do AI gợi ý.'),
    ('MenuPrepOrder / Item', 'meals', 'Đơn nguyên liệu cần chuẩn bị đã xác nhận.'),
    ('MealRegistration / MealTransfer', 'registrations', 'Đăng ký suất ăn và yêu cầu chuyển suất ăn.'),
    ('DailyPurchase / PurchaseExtraItem', 'finance', 'Đơn mua hàng và mặt hàng trong hóa đơn.'),
    ('InventoryEntry / InventoryLog', 'finance', 'Tồn kho và lịch sử nhập - xuất kho.'),
    ('MealPriceSetting', 'finance', 'Cấu hình giá suất ăn theo giai đoạn.'),
    ('DailyNutritionAnalysis', 'finance', 'Kết quả phân tích dinh dưỡng của AI theo ngày.'),
    ('AttendanceLog / AttendanceCapture', 'core', 'Bản ghi điểm danh và ảnh chụp nhận diện.'),
    ('RecognitionHeartbeat / CameraStatusLog', 'core', 'Trạng thái sống và lịch sử online/offline của camera.'),
    ('MealReview / DishReview', 'reviews', 'Đánh giá bữa ăn và chấm điểm từng món.'),
    ('DishSuggestion / Vote', 'reviews', 'Đề xuất món mới và bình chọn.'),
]
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Light Grid Accent 2'
hdr = tbl.rows[0].cells
for i, t in enumerate(['Thực thể (Model)', 'Module', 'Mô tả']):
    hdr[i].paragraphs[0].add_run(t).bold = True
for name, mod, desc in data:
    cells = tbl.add_row().cells
    cells[0].paragraphs[0].add_run(name).bold = True
    cells[1].text = mod
    cells[2].text = desc

# =========================================================
# 7. QUY TRÌNH NGHIỆP VỤ
# =========================================================
h1('7. Quy trình nghiệp vụ tiêu biểu')
h2('7.1. Quy trình một ngày vận hành bếp')
steps = [
    'Bếp trưởng tạo thực đơn ngày (hoặc dùng gợi ý AI) → gửi Admin phê duyệt.',
    'Admin duyệt thực đơn → hệ thống tính nguyên liệu cần mua theo số người đăng ký.',
    'Người ăn đăng ký (import Excel) → Bot NetChat gửi tin nhắn xác nhận / nhắc nhở.',
    'Bếp tạo đơn mua hàng & nhập hóa đơn → Admin duyệt chi phí.',
    'Đến giờ ăn: camera nhận diện điểm danh → đối soát với danh sách đăng ký.',
    'Màn hình TV chiếu thực đơn, chúc mừng sinh nhật; người ăn quét QR đánh giá món.',
    'Cuối ngày: AI phân tích dinh dưỡng, Dashboard tổng hợp thu - chi và chênh lệch.',
]
for i, s in enumerate(steps, 1):
    p = doc.add_paragraph(style='List Number')
    p.add_run(s)

# =========================================================
# 8. TRIỂN KHAI
# =========================================================
h1('8. Triển khai & Vận hành')
bullet('Máy chủ: Ubuntu Server, ứng dụng chạy qua Gunicorn, Nginx làm reverse proxy.')
bullet('Tên miền: net2kitchen.viettel.pro.vn (do Viettel cấp).')
bullet('Cơ sở dữ liệu: PostgreSQL.')
bullet('Các tác vụ định kỳ (cảnh báo camera, hủy chuyển suất quá hạn...) chạy bằng Django management command kết hợp systemd timer.')
bullet('Cấu hình nhạy cảm (Bot, AI key, giá suất ăn) quản lý runtime trong giao diện Admin, không cần sửa code.')

doc.add_paragraph()
end = doc.add_paragraph()
end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run('— Hết —')
r.italic = True
r.font.color.rgb = GREY

out = 'Mo_ta_he_thong_Bep_an.docx'
doc.save(out)
print('Saved:', out)
